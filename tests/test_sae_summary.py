import re
import asyncio
import httpx
import os

from app.main import app


async def _post_file_async(path: str, compound_a: str, compound_b: str):
    async with httpx.AsyncClient(
        transport=httpx.ASGITransport(app=app), base_url="http://test"
    ) as ac:
        with open(path, "rb") as f:
            files = {
                "file": (
                    os.path.basename(path),
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            data = {"compound_a": compound_a, "compound_b": compound_b}
            return await ac.post("/sae-summary", data=data, files=files)


def test_client1_succeeds():
    resp = asyncio.run(_post_file_async("client1_ae.docx", "Placebo", "Compound X"))
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["selected_compounds"] == ["Placebo", "Compound X"]
    assert "tables" in body and isinstance(body["tables"], list)
    # check structure of first table
    t0 = body["tables"][0]
    assert "title" in t0 and "table_number" in t0
    assert "totals_sentence" in t0 and isinstance(t0["totals_sentence"], str)
    assert "sentences" in t0 and isinstance(t0["sentences"], list)


def test_client2_succeeds():
    # client2 file uses different compound column names (e.g., CMP1/CMP2)
    resp = asyncio.run(_post_file_async("client2_ae.docx", "CMP1", "CMP2"))
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["selected_compounds"] == ["CMP1", "CMP2"]
    assert len(body["tables"]) >= 1


def test_sentences_contain_percent_or_zero():
    resp = asyncio.run(_post_file_async("client1_ae.docx", "Placebo", "Compound X"))
    assert resp.status_code == 200, resp.text
    body = resp.json()
    for table in body["tables"]:
        for s in table["sentences"]:
            # sentences should either contain a percent like 5% or state "No participants"
            assert (
                re.search(r"\d+\.?\d*%", s)
                or s.startswith("No participants")
                or s.startswith("An equal proportion")
                or s.startswith("Only ")
            )


def _make_docx_buffer(caption: str, headers: list, rows: list):
    from docx import Document
    import io

    doc = Document()
    if caption:
        doc.add_paragraph(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in rows:
        new_row = table.add_row().cells
        for i, cell in enumerate(r):
            new_row[i].text = str(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def test_non_docx_rejected():
    async def run():
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as ac:
            resp = await ac.post(
                "/sae-summary",
                data={"compound_a": "A", "compound_b": "B"},
                files={"file": ("test.txt", b"not a docx", "text/plain")},
            )
            return resp

    resp = asyncio.run(run())
    assert resp.status_code == 400


def test_missing_compounds_rejected():
    resp = asyncio.run(_post_file_async("client1_ae.docx", "", ""))
    # depending on validation the API may return 400 or 422
    assert resp.status_code in (400, 422)


def test_no_sae_table_returns_400():
    buf = _make_docx_buffer("Not SAE", ["A", "B"], [["1", "(5%)"]])

    async def run():
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as ac:
            resp = await ac.post(
                "/sae-summary",
                data={"compound_a": "A", "compound_b": "B"},
                files={
                    "file": (
                        "nosae.docx",
                        buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            return resp

    resp = asyncio.run(run())
    assert resp.status_code == 400


def test_table_missing_preferred_term_skipped():
    buf = _make_docx_buffer(
        "Table SAE", ["NotPreferred", "Placebo", "Compound X"], [["x", "(1%)", "(2%)"]]
    )

    async def run():
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as ac:
            resp = await ac.post(
                "/sae-summary",
                data={"compound_a": "Placebo", "compound_b": "Compound X"},
                files={
                    "file": (
                        "missing_pref.docx",
                        buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            return resp

    resp = asyncio.run(run())
    # endpoint should not find valid SAE table with preferred term, so 400
    assert resp.status_code == 400


def test_totals_missing_but_sentence_present():
    # create SAE table with preferred term and compound columns but no totals row
    headers = ["Preferred Term", "Placebo", "Compound X"]
    rows = [["Headache", "(0%)", "(5%)"], ["Nausea", "(0%)", "(0%)"]]
    buf = _make_docx_buffer("Serious Adverse Events", headers, rows)

    async def run():
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as ac:
            resp = await ac.post(
                "/sae-summary",
                data={"compound_a": "Placebo", "compound_b": "Compound X"},
                files={
                    "file": (
                        "nototals.docx",
                        buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            return resp

    resp = asyncio.run(run())
    assert resp.status_code == 200
    body = resp.json()
    assert body["tables"][0]["totals_sentence"] == "Totals not found in table."


def test_equal_percentages_and_zero_cases():
    headers = ["Preferred Term", "Placebo", "Compound X"]
    rows = [
        ["Term1", "(5%)", "(5%)"],
        ["Term2", "(0%)", "(0%)"],
        ["Term3", "(0%)", "(2%)"],
    ]
    buf = _make_docx_buffer("Serious AE Table", headers, rows)

    async def run():
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=app), base_url="http://test"
        ) as ac:
            resp = await ac.post(
                "/sae-summary",
                data={"compound_a": "Placebo", "compound_b": "Compound X"},
                files={
                    "file": (
                        "cases.docx",
                        buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            return resp

    resp = asyncio.run(run())
    assert resp.status_code == 200
    body = resp.json()
    sentences = body["tables"][0]["sentences"]
    # expect equal proportion sentence, no participants sentence, and Only xx% sentence
    assert any(
        s.startswith("An equal proportion of participants experienced Term1")
        for s in sentences
    )
    assert any(s.startswith("No participants experienced Term2") for s in sentences)
    assert any(
        ("Only" in s and "Compound X" in s and "Term3" in s and "2" in s)
        for s in sentences
    )
